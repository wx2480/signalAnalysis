import datetime
import os, sys
sys.path.append(r'/data/stock/newSystemData/feature_base/structure')
from DB_Control import db
import logging
from logging import handlers

from tqdm import tqdm
from joblib import Parallel, delayed

import matplotlib.pyplot as plt
from matplotlib.gridspec import GridSpec
import seaborn as sns
import statsmodels.api as sm
import time
import numpy as np
import pandas as pd

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm,Pt
from docx.shared import Inches
from docx.shared import RGBColor

# analysis tools writed by cython
sys.path.append('/home/sharedFold/zhaorui/')
import daily_analysis as da
# obj = da.SignalAnalysis(20200101,20200721,'/home/sharedFold/zhaorui/ret.parquet','close','/home/xiaonan/factor_wxn/factor/bid_ask_spread/')


sys.path.append('/data/stock/newSystemData/feature_base/structure/yili/db_utilise/')
import trading_date as ts
# ts.get_nxt_trading_date(20200721)

sys.path.append('/data/stock/newSystemData/feature_base/structure/yuanyang/alphasystem/dataMgr/')
import dataMgr as dM


def logFactor(filename):
    logger = logging.getLogger('/home/xiaonan/factor_wxn/Log/{}.log'.format(filename))
    logger.setLevel(level = logging.INFO)
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s - %(filename)s - %(funcName)s')

    handler = handlers.TimedRotatingFileHandler('/home/xiaonan/factor_wxn/Log/{}.log'.format(filename), when = 'D', backupCount = 30, encoding = 'utf-8')
    handler.setLevel(logging.INFO)
    handler.setFormatter(formatter)

    console = logging.StreamHandler()
    console.setLevel(logging.INFO)
    console.setFormatter(formatter)

    logger.addHandler(handler)
    logger.addHandler(console)

    return logger


class FactorAnalysis:
    def __init__(self,start,end,factor_list,logger):
        self.start = start
        self.end = end
        self.factor_list = factor_list
        self.ic_values, self.group = {}, {}
        self.logger = logger
        self.pool = dM.load_universe(start,end,univ_name = 'TOP2000')
        self.pool = self.pool.apply(lambda x:x.apply(lambda x:int(x)))

        plt.rcParams['figure.figsize'] = (15,5)
        plt.rcParams['axes.unicode_minus'] = False
        plt.rcParams['xtick.direction'] = 'in'
        plt.rcParams['ytick.direction'] = 'in'

    def get_all_data(self,factor_name):
        timeline = ts.get_trading_date(self.start, self.end)
        ret = pd.read_parquet('/home/sharedFold/zhaorui/ret.parquet')

        read_path = r'/home/xiaonan/factor_wxn/factor/'
        read_path = os.path.join(read_path, factor_name)
        
        ans = []
        
        for date in timeline:
            path = read_path + r'/' + str(date) + '.csv'
            if os.path.exists(path):
                factor = pd.read_csv(path, index_col=None, header=None)
            else:
                continue
            factor.columns = ['code', 'nouse','values']
            factor.loc[:,'date'] = date
            factor = factor[factor['code'].isin(self.pool.loc[date,'code'].values)].copy()
            ans.append(factor)

        factor = pd.concat(ans)
        factor.reset_index(inplace = True)
        factor.drop(['index','nouse'], axis = 1, inplace = True)
        
        
        self.allData = pd.merge(factor, ret, how = 'left', left_on=['code','date'], right_on=['code','dt'])
    
    def get_Ratio(self):
        x = pd.pivot_table(self.allData, index=['code'], columns=['date'], values=['values'])
        
        self.coverRatio = {}
        self.coverRatio[0] = x.apply(lambda x:x.dropna().shape[0]/2000).loc['values']
        self.coverRatio[1] = x.apply(lambda x:x.dropna().drop_duplicates().shape[0]/2000).loc['values']
        _a, a = x.iloc[:,:-1], x.iloc[:,1:]
        _a.columns = a.columns
        self.turnoverRatio = a.corrwith(_a, method = 'spearman').loc['values']
    
    def get_ic(self,factor_name):
        def f(x):
            x.set_index('code',inplace = True)
            ic_0 = x.loc[:,['values']].corrwith(x.loc[:,'y_close_5'], method = 'spearman').iloc[0]    # all ic
            common_index = x.sort_values('values').index
            length = len(common_index)
            ic_1 = x.loc[common_index[:length//2],['values']].corrwith(x.loc[common_index[:length//2],'y_close_5'], method = 'spearman').iloc[0]    # x bottom ic
            ic_2 = x.loc[common_index[length//2:],['values']].corrwith(x.loc[common_index[length//2:],'y_close_5'], method = 'spearman').iloc[0]    # x top ic
            common_index = x.sort_values('y_close_5').index
            length = len(common_index)
            ic_3 = x.loc[common_index[:length//2],['values']].corrwith(x.loc[common_index[:length//2],'y_close_5'], method = 'spearman').iloc[0]    # y bottom ic
            ic_4 = x.loc[common_index[length//2:],['values']].corrwith(x.loc[common_index[length//2:],'y_close_5'], method = 'spearman').iloc[0]    # y top ic

            ans = pd.Series([ic_0, ic_1,ic_2,ic_3,ic_4])
            ans.index = ['AllIC', 'xBottomIC','xTopIC','yBottomIC','yTopIC']

            return ans

        self.ic_values[factor_name] = self.allData.groupby('date').apply(f)
        self.ic_values[factor_name] = self.ic_values[factor_name].sort_index()
        self.ic_values[factor_name].index = self.ic_values[factor_name].index.map(str)

        return self.ic_values
    
    def group_test(self, factor_name, n = 5):
        def f(x):
            ruler = x.loc[:,'values'].drop_duplicates().sort_values().copy()
            try:
                ruler = pd.qcut(ruler, n)
            except:
                print(x.dt.iloc[0])
            ruler = pd.DataFrame(ruler)
            ruler['y'] = x.loc[ruler.index,'values']
            ruler.columns = ['category', 'values']
            x = pd.merge(x, ruler,left_on='values',right_on='values',how='left')
            ans = x.loc[:,['y_close_1','category']].groupby('category').mean()['y_close_1']
            ans.index = ans.index.codes + 1
            return ans.copy()

        self.group[factor_name] = self.allData.groupby('date').apply(f)

        self.group[factor_name] = self.group[factor_name].sort_index()
        self.group[factor_name].index = self.group[factor_name].index.map(str)
        return 1

    def plot(self, factor_name, plot = False):
        save_path = os.path.join(r'/home/xiaonan/factor_wxn/Report/', factor_name)

        fig,ax = plt.subplots(3,1,figsize = (15,15))
        ax[0].bar(x = self.coverRatio[0].index.map(str), height = self.coverRatio[0], color = '#9BBB59')
        ax[0].set_title('coverRatio0  {}'.format(self.coverRatio[0].mean()))
        ax[1].bar(x = self.coverRatio[1].index.map(str), height = self.coverRatio[1], color = '#9BBB59')
        ax[1].set_title('coverRatio1  {}'.format(self.coverRatio[1].mean()))
        ax[2].bar(x = self.turnoverRatio.index.map(str), height = self.turnoverRatio, color = '#9BBB59')
        ax[2].set_title('turnoverRatio  {}'.format(self.turnoverRatio.mean()))
        for i in range(3):
            for j,tick in enumerate(ax[i].get_xticklabels()):
                if j%128 == 0:
                    tick.set_visible(True)
                    tick.set_rotation(30)
                else:
                    tick.set_visible(False)
            for j, tick in enumerate(ax[i].get_xticklines()):
                if j%128 == 0:
                    tick.set_visible(True)
                else:
                    tick.set_visible(False)

        fig.savefig(save_path + '/{}_Ratio.jpg'.format(factor_name))

        fig = plt.figure(figsize = (15, 10), constrained_layout=True)
        gs = GridSpec(3, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,:])
        ax2 = fig.add_subplot(gs[1,0])
        ax3 = fig.add_subplot(gs[1,1])
        ax4 = fig.add_subplot(gs[2,0])
        ax5 = fig.add_subplot(gs[2,1])
        ax = [ax1,ax2,ax3,ax4,ax5]
        labels = ['AllIC','xTopIC','yTopIC','xBottomIC','yBottomIC']
        colors = ['#44A948', '#137CB1', '#EACC80', '#A8D7E2', '#E58061']
        for i in range(5):
            ax[i].plot(self.ic_values[factor_name].loc[:,labels[i]], color = '#00B0F0', label = labels[i])
            ax[i].set_title('{} {}'.format(labels[i], round(self.ic_values[factor_name].loc[:,labels[i]].mean(),4)))
            for j in ['top', 'bottom', 'left', 'right']:
                ax[i].spines[j].set_visible(True)
            ax[i].yaxis.grid(linestyle='--',alpha = 0.3)
            for j,tick in enumerate(ax[i].get_xticklabels()):
                if j%128 == 0:
                    tick.set_visible(True)
                    tick.set_rotation(15)
                else:
                    tick.set_visible(False)
            
            for j,tick in enumerate(ax[i].get_xticklines()):
                if j%128 == 0:
                    tick.set_visible(True)
                else:
                    tick.set_visible(False)
                
        fig.savefig(save_path + '/{}_IC.jpg'.format(factor_name))
        if plot:
            plt.show(fig)

        fig = plt.figure(constrained_layout=True)
        gs = GridSpec(2, 2, figure=fig)
        ax1 = fig.add_subplot(gs[0,:])
        ax2 = fig.add_subplot(gs[1,0])
        ax3 = fig.add_subplot(gs[1,1])
        ax = [ax1,ax2,ax3]
        results = ((self.group[factor_name] + 1) * 0.9995).cumprod()
        
        if isinstance(results, pd.Series):
            results = results.to_frame()
        for i in results.columns:
            ax[0].plot(results.loc[:,i], color = colors[i - 1], label = 'group {}'.format(str(i)))
        ax[0].legend(loc = 'upper left')
        for num in [0,1,2]:
            for i in ['top', 'bottom', 'left', 'right']:
                ax[num].spines[i].set_visible(True)
            ax[num].yaxis.grid(linestyle='--',alpha = 0.3)
        ax[0].set_title('{}'.format(factor_name))

        for i,tick in enumerate(ax[0].get_xticklabels()):
            if i%128 == 0:
                tick.set_visible(True)
                tick.set_rotation(15)
            else:
                tick.set_visible(False)
        for i,tick in enumerate(ax[0].get_xticklines()):
                if j%128 == 0:
                    tick.set_visible(True)
                else:
                    tick.set_visible(False)
        
        annual_ret = np.power(results.iloc[-2,:],256/results.shape[0]) - 1
        ax[1].bar(x=results.columns, height = annual_ret, color = '#00B0F0')

        date = np.random.choice(self.allData['dt'])
        ans = self.allData[self.allData['dt'] == date].copy()
        ax[2].scatter(ans['values'], ans['y_close_1'], color = '#00B0F0',s = 1)
        ax[2].set_title('values and ret1d')
        fig.savefig(save_path + '/{}_Group.jpg'.format(factor_name))

        fig, ax = plt.subplots()
        if self.ic_values[factor_name]['AllIC'].mean() > 0:
            ls = results.loc[:,5] - results.loc[:,1] + 1
        else:
            ls = results.loc[:,1] - results.loc[:,5] + 1
        ax.plot(ls, color = '#00B0F0', label = 'longshort')
        drawback = 1 - ls/np.maximum.accumulate(ls)
        ax1 = ax.twinx()
        ax2 = ax1.twiny()
        ax2.invert_yaxis()
        ax2.stackplot(drawback.index, drawback.values, color = '#A6A6A6', alpha = 0.3)
        
        ax.yaxis.grid(linestyle='--',alpha = 0.3)

        for i,tick in enumerate(ax.get_xticklabels()):
            if i%128 == 0:
                tick.set_visible(True)
                tick.set_rotation(15)
            else:
                tick.set_visible(False)
        
        for i,tick in enumerate(ax.get_xticklines()):
            if i%128 == 0:
                tick.set_visible(True)
            else:
                tick.set_visible(False)
        for i,tick in enumerate(ax2.get_xticklabels()):
            tick.set_visible(False)
        for i,tick in enumerate(ax2.get_xticklines()):
            tick.set_visible(False)
        
        plt.legend('upper left')
        fig.savefig(save_path + '/{}_ls.jpg'.format(factor_name))
        if plot:
            plt.show(fig)
        

    def analysis(self):
        for factor_name in self.factor_list:
            if os.path.exists(os.path.join(r'/home/xiaonan/factor_wxn/Report/', factor_name)):
                pass
            else:
                os.mkdir(os.path.join(r'/home/xiaonan/factor_wxn/Report/', factor_name))
            self.logger.info('{} start.'.format(factor_name))
            try:
                self.get_all_data(factor_name)
                self.get_Ratio()
                self.get_ic(factor_name)
                self.group_test(factor_name)
                self.plot(factor_name, plot = False)
                self.output_report(factor_name)
            except:
                self.logger.info('{} FAILED!!!.'.format(factor_name))
                continue
            



    def output_report(self, factor_name):
        path = r'/home/xiaonan/factor_wxn/Report/'

        _path = os.path.join(path, factor_name)
        if os.path.exists(_path):
            pass
        else:
            os.mkdir(_path)
        
        report = Document()
        styles = report.styles['Normal']

        title = report.add_heading()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('Factor_Ratio')
        title_run.bold = True
        title_run.font.size = Pt(32)

        report.add_picture(r'/home/xiaonan/factor_wxn/Report/{}/{}_Ratio.jpg'.format(factor_name, factor_name),width = Inches(5.5))

        title = report.add_heading()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('IC')
        title_run.bold = True
        title_run.font.size = Pt(32)

        report.add_picture(r'/home/xiaonan/factor_wxn/Report/{}/{}_IC.jpg'.format(factor_name, factor_name),width = Inches(5.5))

        title = report.add_heading()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('Group')
        title_run.bold = True
        title_run.font.size = Pt(32)

        report.add_picture(r'/home/xiaonan/factor_wxn/Report/{}/{}_Group.jpg'.format(factor_name, factor_name),width = Inches(5.5))

        title = report.add_heading()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('longShort')
        title_run.bold = True
        title_run.font.size = Pt(32)

        report.add_picture(r'/home/xiaonan/factor_wxn/Report/{}/{}_ls.jpg'.format(factor_name, factor_name),width = Inches(5.5))

        results = ((self.group[factor_name] + 1) * 0.9995).cumprod()
        if self.ic_values[factor_name]['AllIC'].mean() > 0:
            ls = results.loc[:,5] - results.loc[:,1] + 1
            lRatio = (results.loc[:,5].diff()/ls.diff()).mean()
            sRatio = (-results.loc[:,1].diff()/ls.diff()).mean()
            sharpe = (np.power(ls.iloc[-2], 256/ls.shape[0]) - 1 - 0.03)/(ls.diff().std() * 16)
        else:
            ls = results.loc[:,1] - results.loc[:,5] + 1
            lRatio = (results.loc[:,1].diff()/ls.diff()).mean()
            sRatio = (-results.loc[:,5].diff()/ls.diff()).mean()
            sharpe = (np.power(ls.iloc[-2], 256/ls.shape[0]) - 1 - 0.03)/(ls.diff().std() * 16)
        
        report.add_paragraph('\n')
        table = report.add_table(rows = 2, cols = 3)        
        p = table.rows[0].cells[0].add_paragraph('longRatio')
        p = table.rows[1].cells[0].add_paragraph(str(round(lRatio, 4)))
        p = table.rows[0].cells[1].add_paragraph('shortRatio')
        p = table.rows[1].cells[1].add_paragraph(str(round(sRatio, 4)))
        p = table.rows[0].cells[2].add_paragraph('Sharpe')
        p = table.rows[1].cells[2].add_paragraph(str(round(sharpe, 4)))

        
        report.save(r'/home/xiaonan/factor_wxn/Report/{}/{}.docx'.format(factor_name, factor_name))

if __name__ == '__main__':
    logger = logFactor('analysis')
    logger.info('analysis start.')
    start = 20150106
    end = 20200801
    a = datetime.datetime.now()
    # factor_name = 'askid_num'
    factor_name =  'price_vol_corr'
    factor_list = os.listdir('/home/xiaonan/factor_wxn/factor/')
    out_list = ['corr_var','corr_mean','corr_skew','corr_kurt','liquid_shock_elasticity','price2vol_shock','withdraw_ratio','pareto_buy','open_auction_netinflows','enormous_bprice_location','enormous_sprice_location','vol_mean_price','true_enormous_order']
    for i in out_list:
        if i in factor_list:
            factor_list.remove(i)
    if '.directort' in factor_list:
        factor_list.remove('.directory')
    '''
    factor_list = [
        #'withdraw_ratio',
        'pareto_buy','open_auction_netinflows','enormous_bprice_location','enormous_sprice_location','vol_mean_price', 'price2vol_shock']
    '''
    factor_list = os.listdir('/home/xiaonan/factor_wxn/factor/')
    A = FactorAnalysis(start, end, factor_list, logger)
    A.analysis()

    b = datetime.datetime.now()
    print('')
    print(b-a)
    print('')
    logger.info('analysis end.')
    '''
    output_report(['market_depth_amount'])
    '''
